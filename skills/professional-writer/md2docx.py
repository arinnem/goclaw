"""
Markdown to DOCX Converter
===========================
Converts markdown files to well-formatted Word documents,
preserving headings, tables, lists, code blocks, blockquotes,
inline formatting (bold, italic, code, links), and horizontal rules.

Usage:
    python md2docx.py input.md [output.docx]
    python md2docx.py input.md --font "Times New Roman" --font-size 12

Requirements:
    pip install python-docx markdown beautifulsoup4
"""

import sys
import argparse
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


class MarkdownToDocxConverter:
    """Converts Markdown files to well-formatted DOCX documents."""

    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(12)
    HEADING_FONT = "Times New Roman"
    HEADING_COLOR = RGBColor(0x1A, 0x1A, 0x2E)
    HEADING_SIZES = {1: Pt(20), 2: Pt(18), 3: Pt(16), 4: Pt(14), 5: Pt(12), 6: Pt(11)}

    CODE_FONT = "Consolas"
    CODE_FONT_SIZE = Pt(9)
    CODE_BG = "F5F5F5"
    CODE_BORDER = "D0D0D0"

    BLOCKQUOTE_TEXT_COLOR = RGBColor(0x6A, 0x73, 0x7D)
    BLOCKQUOTE_BORDER_COLOR = "D0D7DE"

    TABLE_HEADER_BG = "E2E8F0"
    TABLE_FONT_SIZE = Pt(10)

    LINK_COLOR = RGBColor(0x06, 0x69, 0xAD)

    MD_EXTENSIONS = ["tables", "fenced_code", "sane_lists", "smarty", "attr_list"]

    def __init__(self, input_path: str, output_path: str = None):
        self.input_path = Path(input_path)
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {self.input_path}")
        
        base_output_path = Path(output_path) if output_path else self.input_path.with_suffix(".docx")
        
        # Versioning logic to avoid overwriting existing files
        if base_output_path.exists():
            version = 1
            while True:
                new_path = base_output_path.with_name(f"{base_output_path.stem}-v{version}{base_output_path.suffix}")
                if not new_path.exists():
                    self.output_path = new_path
                    break
                version += 1
        else:
            self.output_path = base_output_path

        self.doc = Document()
        self._list_depth = 0

    def convert(self) -> Path:
        md_text = self.input_path.read_text(encoding="utf-8")
        html = markdown.markdown(md_text, extensions=self.MD_EXTENSIONS)
        soup = BeautifulSoup(html, "html.parser")

        self._setup_document()
        self._process_children(soup)

        self.doc.save(str(self.output_path))
        return self.output_path

    # ── Document setup ──────────────────────────────────────────────

    def _setup_document(self):
        style = self.doc.styles["Normal"]
        style.font.name = self.FONT_NAME
        style.font.size = self.FONT_SIZE

        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.15

        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    # ── Tree walker ─────────────────────────────────────────────────

    def _process_children(self, element):
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p = self.doc.add_paragraph()
                    p.add_run(text)
            elif isinstance(child, Tag):
                self._process_element(child)

    def _process_element(self, el: Tag):
        tag = el.name
        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            self._add_heading(el, int(tag[1]))
        elif tag == "p":
            self._add_paragraph(el)
        elif tag == "ul":
            self._add_list(el, ordered=False)
        elif tag == "ol":
            self._add_list(el, ordered=True)
        elif tag == "table":
            self._add_table(el)
        elif tag == "pre":
            self._add_code_block(el)
        elif tag == "blockquote":
            self._add_blockquote(el)
        elif tag == "hr":
            self._add_horizontal_rule()
        elif tag in ("div", "section", "article", "main"):
            self._process_children(el)
        elif tag == "br":
            pass
        else:
            self._add_paragraph(el)

    # ── Headings ────────────────────────────────────────────────────

    def _add_heading(self, el: Tag, level: int):
        heading = self.doc.add_heading(level=level)
        self._add_inline(el, heading)

        for run in heading.runs:
            run.font.name = self.HEADING_FONT
            run.font.size = self.HEADING_SIZES.get(level, Pt(11))
            run.font.color.rgb = self.HEADING_COLOR
            # East-Asian font fallback for Vietnamese diacritics
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{self.HEADING_FONT}"/>')
                rpr.insert(0, rFonts)
            else:
                rFonts.set(qn("w:eastAsia"), self.HEADING_FONT)

        heading.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        heading.paragraph_format.space_after = Pt(8)

    # ── Paragraphs ──────────────────────────────────────────────────

    def _add_paragraph(self, el: Tag):
        p = self.doc.add_paragraph()
        self._add_inline(el, p)

    # ── Inline content (recursive) ──────────────────────────────────

    def _add_inline(self, el, para, *, bold=False, italic=False, code=False, strike=False):
        for child in el.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if not text:
                    continue
                run = para.add_run(text)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                if strike:
                    run.font.strike = True
                if code:
                    run.font.name = self.CODE_FONT
                    run.font.size = self.CODE_FONT_SIZE
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.CODE_BG}"/>')
                    run._element.get_or_add_rPr().append(shd)

            elif isinstance(child, Tag):
                t = child.name
                if t in ("strong", "b"):
                    self._add_inline(child, para, bold=True, italic=italic, code=code, strike=strike)
                elif t in ("em", "i"):
                    self._add_inline(child, para, bold=bold, italic=True, code=code, strike=strike)
                elif t == "code":
                    self._add_inline(child, para, bold=bold, italic=italic, code=True, strike=strike)
                elif t in ("del", "s", "strike"):
                    self._add_inline(child, para, bold=bold, italic=italic, code=code, strike=True)
                elif t == "a":
                    self._add_hyperlink(child, para)
                elif t == "br":
                    para.add_run("\n")
                elif t == "img":
                    self._add_image(child, para)
                elif t == "sup":
                    r = para.add_run(child.get_text())
                    r.font.superscript = True
                elif t == "sub":
                    r = para.add_run(child.get_text())
                    r.font.subscript = True
                else:
                    self._add_inline(child, para, bold=bold, italic=italic, code=code, strike=strike)

    # ── Hyperlinks ──────────────────────────────────────────────────

    def _add_hyperlink(self, el: Tag, para):
        url = el.get("href", "")
        text = el.get_text()

        part = para.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        )
        new_run = parse_xml(
            f'<w:r {nsdecls("w")}>'
            f"<w:rPr>"
            f'<w:color w:val="0669AD"/>'
            f'<w:u w:val="single"/>'
            f"</w:rPr>"
            f'<w:t xml:space="preserve">{_escape_xml(text)}</w:t>'
            f"</w:r>"
        )
        hyperlink.append(new_run)
        para._element.append(hyperlink)

    # ── Images ──────────────────────────────────────────────────────

    def _add_image(self, el: Tag, para):
        src = el.get("src", "")
        alt = el.get("alt", "")
        img_path = Path(src)
        if not img_path.is_absolute():
            img_path = self.input_path.parent / src

        if img_path.exists():
            try:
                run = para.add_run()
                run.add_picture(str(img_path), width=Inches(5))
                return
            except Exception:
                pass

        r = para.add_run(f"[Image: {alt or src}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Lists ───────────────────────────────────────────────────────

    def _add_list(self, el: Tag, ordered=False):
        bullet_chars = ["•", "◦", "▪", "▹"]
        for idx, li in enumerate(el.find_all("li", recursive=False)):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            indent = self._list_depth + 1
            pf.left_indent = Cm(1.27 * indent)
            pf.first_line_indent = Cm(-0.63)
            pf.space_after = Pt(2)
            pf.space_before = Pt(1)

            prefix = f"{idx + 1}. " if ordered else f"{bullet_chars[min(self._list_depth, 3)]} "
            run = p.add_run(prefix)
            if not ordered:
                run.font.size = Pt(10)

            # Process inline content of this <li>, skipping nested lists
            for child in li.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        p.add_run(text)
                elif isinstance(child, Tag):
                    if child.name in ("ul", "ol"):
                        self._list_depth += 1
                        self._add_list(child, ordered=(child.name == "ol"))
                        self._list_depth -= 1
                    elif child.name == "p":
                        self._add_inline(child, p)
                    else:
                        self._add_inline_single(child, p)

    def _add_inline_single(self, el: Tag, para):
        """Handle a single inline tag within a list item."""
        t = el.name
        if t in ("strong", "b"):
            self._add_inline(el, para, bold=True)
        elif t in ("em", "i"):
            self._add_inline(el, para, italic=True)
        elif t == "code":
            self._add_inline(el, para, code=True)
        elif t == "a":
            self._add_hyperlink(el, para)
        else:
            self._add_inline(el, para)

    # ── Tables ──────────────────────────────────────────────────────

    def _add_table(self, el: Tag):
        header_data, body_data = [], []

        thead = el.find("thead")
        if thead:
            for tr in thead.find_all("tr"):
                header_data.append(self._extract_row_cells(tr))

        tbody = el.find("tbody")
        rows_source = tbody if tbody else el
        for tr in rows_source.find_all("tr", recursive=(tbody is None)):
            if thead and tr.parent == thead:
                continue
            cells = self._extract_row_cells(tr)
            if cells:
                body_data.append(cells)

        all_rows = header_data + body_data
        if not all_rows:
            return

        n_cols = max(len(r) for r in all_rows)
        n_rows = len(all_rows)

        # Add a small spacer before the table
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(0)

        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Apply Table Grid style for borders
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for i, row_data in enumerate(all_rows):
            row = table.rows[i]
            for j, cell_content in enumerate(row_data):
                if j >= n_cols:
                    break
                cell = row.cells[j]

                # Clear default paragraph and write content
                cell.text = ""
                p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                p.text = cell_content

                for run in p.runs:
                    run.font.name = self.FONT_NAME
                    run.font.size = self.TABLE_FONT_SIZE

                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)

                # Style header rows
                if i < len(header_data):
                    for run in p.runs:
                        run.bold = True
                    shd = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{self.TABLE_HEADER_BG}" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shd)

        # Spacer after table
        spacer2 = self.doc.add_paragraph()
        spacer2.paragraph_format.space_before = Pt(0)
        spacer2.paragraph_format.space_after = Pt(4)

    def _extract_row_cells(self, tr: Tag) -> list[str]:
        return [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]

    # ── Code blocks ─────────────────────────────────────────────────

    def _add_code_block(self, el: Tag):
        code_el = el.find("code")
        code_text = code_el.get_text() if code_el else el.get_text()

        # Remove trailing newline if present
        if code_text.endswith("\n"):
            code_text = code_text[:-1]

        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after = Pt(8)
        pf.left_indent = Cm(0.5)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # Background shading
        pPr = p._element.get_or_add_pPr()
        pPr.append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{self.CODE_BG}" w:val="clear"/>')
        )
        # Border
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:top w:val="single" w:sz="4" w:space="4" w:color="{self.CODE_BORDER}"/>'
                f'<w:left w:val="single" w:sz="4" w:space="4" w:color="{self.CODE_BORDER}"/>'
                f'<w:bottom w:val="single" w:sz="4" w:space="4" w:color="{self.CODE_BORDER}"/>'
                f'<w:right w:val="single" w:sz="4" w:space="4" w:color="{self.CODE_BORDER}"/>'
                f"</w:pBdr>"
            )
        )

        lines = code_text.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run("\n")
            run = p.add_run(line)
            run.font.name = self.CODE_FONT
            run.font.size = self.CODE_FONT_SIZE
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Blockquotes ─────────────────────────────────────────────────

    def _add_blockquote(self, el: Tag):
        for child in el.children:
            if isinstance(child, Tag):
                if child.name == "p":
                    p = self.doc.add_paragraph()
                    pf = p.paragraph_format
                    pf.left_indent = Cm(1.5)
                    pf.space_before = Pt(4)
                    pf.space_after = Pt(4)

                    # Left border for blockquote
                    pPr = p._element.get_or_add_pPr()
                    pPr.append(
                        parse_xml(
                            f'<w:pBdr {nsdecls("w")}>'
                            f'<w:left w:val="single" w:sz="18" w:space="8" w:color="{self.BLOCKQUOTE_BORDER_COLOR}"/>'
                            f"</w:pBdr>"
                        )
                    )
                    self._add_inline(child, p)
                    for run in p.runs:
                        if not run.font.color.rgb:
                            run.font.color.rgb = self.BLOCKQUOTE_TEXT_COLOR
                elif child.name == "blockquote":
                    self._add_blockquote(child)
                else:
                    self._process_element(child)
            elif isinstance(child, NavigableString):
                text = str(child).strip()
                if text:
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.5)
                    run = p.add_run(text)
                    run.font.color.rgb = self.BLOCKQUOTE_TEXT_COLOR

    # ── Horizontal rule ─────────────────────────────────────────────

    def _add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._element.get_or_add_pPr()
        pPr.append(
            parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="CCCCCC"/>'
                f"</w:pBdr>"
            )
        )


# ── Helpers ─────────────────────────────────────────────────────────


def _escape_xml(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&apos;")
    )


# ── CLI ─────────────────────────────────────────────────────────────


def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown files to well-formatted DOCX documents.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python md2docx.py document.md
  python md2docx.py document.md output.docx
  python md2docx.py document.md --font "Times New Roman" --font-size 12
  python md2docx.py document.md --code-font "Courier New"
        """,
    )
    parser.add_argument("input", help="Input Markdown file path")
    parser.add_argument(
        "output", nargs="?", help="Output DOCX file path (default: same name .docx)"
    )
    parser.add_argument("--font", default="Calibri", help="Base font (default: Calibri)")
    parser.add_argument(
        "--font-size", type=float, default=11, help="Base font size in pt (default: 11)"
    )
    parser.add_argument(
        "--code-font", default="Consolas", help="Code font (default: Consolas)"
    )

    args = parser.parse_args()

    converter = MarkdownToDocxConverter(args.input, args.output)
    converter.FONT_NAME = args.font
    converter.HEADING_FONT = args.font
    converter.FONT_SIZE = Pt(args.font_size)
    converter.CODE_FONT = args.code_font

    output = converter.convert()
    print(f"✅ Converted: {output}")


if __name__ == "__main__":
    main()
